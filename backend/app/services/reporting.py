import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.schemas.reporting import AnalysisReportRequest


GREEN = RGBColor(29, 106, 71)
DARK_GREEN = RGBColor(36, 75, 57)
MUTED = RGBColor(102, 114, 108)
LIGHT_FILL = "EEF4F0"
GRAY_FILL = "F2F4F3"
DARK_GREEN_FILL = "244B39"
TABLE_WIDTH_DXA = 9360


def _font(run, size: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_margins(cell)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    _font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    tail = paragraph.add_run(" 页")
    _font(tail, 9, color=MUTED)


def _format_value(value: float, metric_type: str) -> str:
    return f"{value * 100:.2f}%" if metric_type == "proportion" else f"{value:.3f}"


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _font(run, 16 if level == 1 else 13, True, GREEN if level == 1 else DARK_GREEN)


def build_analysis_report(payload: AnalysisReportRequest) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = GREEN if name == "Heading 1" else DARK_GREEN
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "CausaLab · 智能实验设计与效应评估平台"
    _font(header.runs[0], 9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_field(footer)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    _font(kicker.add_run("EXPERIMENT ANALYSIS REPORT"), 9, True, GREEN)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _font(title.add_run(payload.project_name), 24, True, DARK_GREEN)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    _font(subtitle.add_run("A/B 实验效应评估报告"), 14, color=MUTED)

    metadata = document.add_table(rows=4, cols=2)
    metadata.style = "Table Grid"
    metadata_data = [
        ("报告生成", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("分析人员", payload.analyst),
        ("核心指标", payload.metric_name),
        ("统计方法", payload.method),
    ]
    for index, (label, value) in enumerate(metadata_data):
        metadata.cell(index, 0).text = label
        metadata.cell(index, 1).text = value
        _shade(metadata.cell(index, 0), GRAY_FILL)
        _font(metadata.cell(index, 0).paragraphs[0].runs[0], 10, True, DARK_GREEN)
        _font(metadata.cell(index, 1).paragraphs[0].runs[0], 10)
    _set_table_geometry(metadata, [2700, 6660])

    _add_heading(document, "1. 执行摘要")
    callout = document.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    _shade(callout.cell(0, 0), LIGHT_FILL)
    callout.cell(0, 0).text = f"统计决策：{payload.decision}\n{payload.conclusion}"
    for run in callout.cell(0, 0).paragraphs[0].runs:
        _font(run, 11, True if run.text.startswith("统计决策") else False, DARK_GREEN)
    _set_table_geometry(callout, [TABLE_WIDTH_DXA])

    _add_heading(document, "2. 核心结果")
    result_table = document.add_table(rows=2, cols=4)
    result_table.style = "Table Grid"
    headers = [payload.control_label, payload.treatment_label, "绝对效应", "p 值"]
    values = [
        f"{_format_value(payload.control_value, payload.metric_type)}\nn = {payload.control_size}",
        f"{_format_value(payload.treatment_value, payload.metric_type)}\nn = {payload.treatment_size}",
        _format_value(payload.absolute_effect, payload.metric_type),
        "< 0.0001" if payload.p_value < 0.0001 else f"{payload.p_value:.4f}",
    ]
    for index in range(4):
        result_table.cell(0, index).text = headers[index]
        result_table.cell(1, index).text = values[index]
        _shade(result_table.cell(0, index), DARK_GREEN_FILL)
        for run in result_table.cell(0, index).paragraphs[0].runs:
            _font(run, 10, True, RGBColor(255, 255, 255))
        for run in result_table.cell(1, index).paragraphs[0].runs:
            _font(run, 11, True if index >= 2 else False, GREEN if index >= 2 else None)
        result_table.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        result_table.cell(1, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_table_geometry(result_table, [2400, 2400, 2280, 2280])

    interval = document.add_paragraph()
    interval.paragraph_format.space_before = Pt(8)
    _font(interval.add_run(f"{payload.confidence_level * 100:.0f}% 置信区间："), 11, True, DARK_GREEN)
    _font(interval.add_run(f"[{_format_value(payload.ci_lower, payload.metric_type)}, {_format_value(payload.ci_upper, payload.metric_type)}]"), 11)
    if payload.relative_effect_percent is not None:
        relative = document.add_paragraph()
        _font(relative.add_run("相对变化："), 11, True, DARK_GREEN)
        _font(relative.add_run(f"{payload.relative_effect_percent:.2f}%"), 11)

    _add_heading(document, "3. 方法与判断规则")
    method_text = (
        f"本报告使用 {payload.method}。显著性水平 α = {payload.alpha:.3f}，"
        f"当 p 值小于 α 时拒绝原假设。报告同时给出效应大小与置信区间，"
        "以避免仅依据统计显著性判断业务价值。"
    )
    document.add_paragraph(method_text)
    if payload.quality_score is not None:
        document.add_paragraph(f"分析前数据质量评分为 {payload.quality_score}/100；该评分应与缺失、重复、异常值及分组比例诊断结合解释。")

    _add_heading(document, "4. 结论与使用建议")
    document.add_paragraph(payload.conclusion)
    for note in payload.notes or ["统计结果反映当前样本证据，不等同于业务收益或长期因果效应。", "上线决策还应结合实验功效、数据质量和实施成本。"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        _font(paragraph.add_run(note), 11)

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(18)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(closing.add_run("本报告由 CausaLab V1.0 自动生成"), 9, color=MUTED)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
